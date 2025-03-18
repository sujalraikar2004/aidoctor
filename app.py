import logging
import os
import platform
import subprocess
import datetime
import base64
import gradio as gr
from io import BytesIO
from dotenv import load_dotenv
from langdetect import detect
from speech_recognition import Recognizer, Microphone
from pydub import AudioSegment
from gtts import gTTS
from groq import Groq
from docx import Document

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
STT_MODEL = "whisper-large-v3"
TEXT_MODEL = "llama3-8b-8192"
VISION_MODEL = "llama-3.2-11b-vision-preview"

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Language map for Groq and gTTS
LANGUAGE_MAP = {
    "English": "en",
    "Hindi": "hi",
    "Marathi": "mr",
    "Kannada": "kn"
}

# ---- AUDIO RECORDING ----
def record_audio(file_path, timeout=20, phrase_time_limit=None):
    recognizer = Recognizer()
    try:
        with Microphone() as source:
            logging.info("Adjusting for ambient noise...")
            recognizer.adjust_for_ambient_noise(source, duration=1)
            logging.info("Start speaking now...")
            audio_data = recognizer.listen(source, timeout=timeout, phrase_time_limit=phrase_time_limit)
            logging.info("Recording complete.")
            
            wav_data = audio_data.get_wav_data()
            audio_segment = AudioSegment.from_wav(BytesIO(wav_data))
            audio_segment.export(file_path, format="mp3", bitrate="128k")
            logging.info(f"Audio saved to {file_path}")
    except Exception as e:
        logging.error(f"An error occurred: {e}")

# ---- SPEECH TO TEXT ----
def transcribe_audio(file_path, language):
    try:
        client = Groq(api_key=GROQ_API_KEY)
        with open(file_path, "rb") as audio_file:
            result = client.audio.transcriptions.create(
                file=audio_file, model=STT_MODEL, language=language
            )
        return result.text
    except Exception as e:
        logging.error(f"Transcription failed: {e}")
        return "Could not process audio input"

# ---- TEXT ANALYSIS ----
def analyze_text(query, language):
    try:
        client = Groq(api_key=GROQ_API_KEY)
        messages = [
            {"role": "system", "content": f"You are a medical assistant providing professional responses in {language}."},
            {"role": "user", "content": query}
        ]
        response = client.chat.completions.create(
            model=TEXT_MODEL, messages=messages, temperature=0.2, max_tokens=500
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"Text analysis error: {e}")
        return "Could not process query."

# ---- TEXT TO SPEECH ----
def text_to_speech(text, language, output_filepath="doctor_response"):
    try:
        mp3_path = f"{output_filepath}.mp3"
        tts = gTTS(text=text, lang=language)
        tts.save(mp3_path)
        return mp3_path
    except Exception as e:
        logging.error(f"TTS error: {e}")
        return None

# ---- IMAGE PROCESSING ----
def encode_image(image_path):
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except Exception as e:
        logging.error(f"Image encoding failed: {e}")
        return ""

def analyze_image(query, image_path):
    try:
        encoded_image = encode_image(image_path)
        client = Groq(api_key=GROQ_API_KEY)
        messages = [{
            "role": "user", "content": [
                {"type": "text", "text": query},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded_image}"}}
            ]
        }]
        response = client.chat.completions.create(
            model=VISION_MODEL, messages=messages, temperature=0.2, max_tokens=500
        )
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"Vision API error: {e}")
        return "Could not analyze image."

# ---- GENERATE PRESCRIPTION ----
def create_prescription(patient_query, doctor_response):
    try:
        os.makedirs("prescriptions", exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"prescriptions/prescription_{timestamp}.docx"
        doc = Document()
        doc.add_heading('AI Medical Prescription', 0)
        doc.add_heading('Patient Concerns', level=2)
        doc.add_paragraph(patient_query)
        doc.add_heading('Medical Analysis', level=1)
        doc.add_paragraph(doctor_response)
        doc.add_paragraph("\nThis prescription should be verified by a licensed medical professional.")
        doc.save(filename)
        return filename
    except Exception as e:
        logging.error(f"Error creating prescription: {e}")
        return None

# ---- GRADIO UI ----
def process_inputs(audio_path, text_query, image_path, language):
    lang_code = LANGUAGE_MAP.get(language, "en")
    patient_query = ""
    
    if audio_path:
        patient_query = transcribe_audio(audio_path, lang_code)
    elif text_query:
        patient_query = text_query
    
    if image_path:
        doctor_response = analyze_image(patient_query, image_path)
    else:
        doctor_response = analyze_text(patient_query, lang_code)
    
    audio_output = text_to_speech(doctor_response, lang_code)
    prescription_path = create_prescription(patient_query, doctor_response)
    
    return [patient_query, doctor_response, audio_output, prescription_path]

iface = gr.Interface(
    fn=process_inputs,
    inputs=[
        gr.Audio(sources=["microphone"], type="filepath", label="Voice Input"),
        gr.Textbox(label="Text Input", placeholder="Type your query here..."),
        gr.Image(type="filepath", label="Skin Condition Photo"),
        gr.Dropdown(choices=["English", "Hindi", "Marathi", "Kannada"], label="Select Language", value="English")
    ],
    outputs=[
        gr.Textbox(label="Patient Query"),
        gr.Textbox(label="Medical Analysis"),
        gr.Audio(label="Voice Response", autoplay=True),
        gr.File(label="Download Prescription")
    ],
    title="AI Smart Medical Assistant",
    description="Supports English, Hindi, Marathi, Kannada. Upload an image for analysis.",
    allow_flagging="never"
)

if __name__ == "__main__":
    iface.launch(server_port=7860, show_error=True)
